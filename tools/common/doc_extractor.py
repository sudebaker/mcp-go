#!/usr/bin/env python3
"""
Document text extraction utilities for MCP Python tools.

Provides functions to extract text from various document formats:
- PDF (.pdf) using pypdf
- Word documents (.docx) using python-docx
- Plain text (.txt)
- Markdown (.md)
- YAML (.yaml, .yml)
- JSON (.json)
- TOML (.toml)
- INI/ENV files (.ini, .env)
"""

import os
import io
import json
import yaml
import configparser
import ipaddress
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Optional
from urllib.parse import urlparse

try:
    from langchain_community.document_loaders import PyPDFLoader

    PYPDF_AVAILABLE = True
except ImportError:
    try:
        from pypdf import PdfReader

        PYPDF_AVAILABLE = True
    except ImportError:
        PYPDF_AVAILABLE = False
        PdfReader = None

try:
    from docx import Document

    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

try:
    import httpx

    HTTPX_AVAILABLE = True
except ImportError:
    httpx = None
    HTTPX_AVAILABLE = False


@dataclass
class ExtractionResult:
    """Result of document text extraction."""

    text: str
    page_count: int = 0
    file_type: str = "unknown"
    metadata: dict = field(default_factory=dict)


def validate_url_for_download(url: str) -> tuple[bool, str]:
    """
    Validate URL to prevent SSRF attacks.
    
    Security: Blocks private/reserved IP ranges and dangerous protocols.
    
    Args:
        url: URL to validate
        
    Returns:
        Tuple of (is_valid, reason_if_invalid)
    """
    try:
        parsed = urlparse(url)
        
        # Only allow HTTP and HTTPS
        if parsed.scheme not in ('http', 'https'):
            return False, f"Protocol '{parsed.scheme}' not allowed. Only HTTP and HTTPS permitted."
        
        # Extract hostname
        hostname = parsed.hostname
        if not hostname:
            return False, "URL must have a valid hostname"
        
        # Validate hostname is not a known cloud metadata service
        metadata_services = [
            'metadata.google.internal',
            '169.254.169.254',
            'metadata',
            'localhost',
            '127.0.0.1',
            '::1',
        ]
        if hostname in metadata_services:
            return False, f"Access to '{hostname}' is not permitted"
        
        # Try to parse as IP address
        try:
            ip = ipaddress.ip_address(hostname)
            # Block private/reserved ranges
            if ip.is_private:
                return False, f"Private IP range blocked: {hostname}"
            if ip.is_loopback:
                return False, f"Loopback IP blocked: {hostname}"
            if ip.is_link_local:
                return False, f"Link-local IP blocked: {hostname}"
            if ip.is_multicast:
                return False, f"Multicast IP blocked: {hostname}"
            if ip.is_reserved:
                return False, f"Reserved IP blocked: {hostname}"
        except ValueError:
            # Not an IP, so it's a domain name - allow if it resolves
            pass
        
        # Validate port is reasonable
        if parsed.port:
            if parsed.port < 1 or parsed.port > 65535:
                return False, f"Invalid port: {parsed.port}"
            # Block common internal service ports
            internal_ports = [
                22,    # SSH
                3306,  # MySQL
                5432,  # PostgreSQL
                6379,  # Redis
                27017, # MongoDB
                9200,  # Elasticsearch
            ]
            if parsed.port in internal_ports:
                return False, f"Port {parsed.port} is blocked for security"
        
        return True, ""
        
    except Exception as e:
        return False, f"URL validation error: {str(e)}"


def download_file(url: str) -> io.BytesIO:
    """
    Download a file from URL and return as BytesIO buffer.

    Args:
        url: URL to download from

    Returns:
        BytesIO buffer containing the file content

    Raises:
        ImportError: If httpx is not available
        ValueError: If URL fails SSRF validation
        Exception: If download fails
    """
    if httpx is None:
        raise ImportError("httpx is not installed. Install with: pip install httpx")

    # SECURITY: Validate URL to prevent SSRF attacks
    is_valid, error_reason = validate_url_for_download(url)
    if not is_valid:
        raise ValueError(f"URL validation failed: {error_reason}")

    try:
        with httpx.Client(timeout=60.0) as client:
            # CRITICAL: For S3 presigned URLs, do NOT follow redirects (breaks signature)
            # Presigned URLs authenticate via query parameters (X-Amz-*) which are
            # invalidated by redirect chains. Also avoid adding extra headers.
            response = client.get(url, follow_redirects=False)
            response.raise_for_status()
            return io.BytesIO(response.content)
    except Exception as e:
        raise Exception(f"Failed to download file from URL: {str(e)}") from e


def extract_text_from_pdf(buffer: io.BytesIO) -> ExtractionResult:
    """
    Extract text from a PDF buffer.

    Args:
        buffer: BytesIO buffer containing PDF content

    Returns:
        ExtractionResult with extracted text and metadata

    Raises:
        ImportError: If pypdf is not available
    """
    if PdfReader is None:
        raise ImportError("pypdf is not installed. Install with: pip install pypdf")

    buffer.seek(0)

    try:
        reader = PdfReader(buffer)
        pages = reader.pages
        text_parts = []
        metadata = {}

        if reader.metadata:
            metadata = {
                k.strip("/"): str(v)
                for k, v in reader.metadata.items()
                if v
            }

        for page in pages:
            text_parts.append(page.extract_text())

        text = "\n".join(text_parts)

        return ExtractionResult(
            text=text,
            page_count=len(pages),
            file_type="pdf",
            metadata=metadata,
        )
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}") from e


def extract_text_from_docx(buffer: io.BytesIO) -> ExtractionResult:
    """
    Extract text from a DOCX buffer.

    Args:
        buffer: BytesIO buffer containing DOCX content

    Returns:
        ExtractionResult with extracted text and metadata

    Raises:
        ImportError: If python-docx is not available
    """
    if Document is None:
        raise ImportError(
            "python-docx is not installed. Install with: pip install python-docx"
        )

    buffer.seek(0)

    try:
        doc = Document(buffer)
        text_parts = []
        metadata = {}

        if doc.core_properties.title:
            metadata["title"] = doc.core_properties.title
        if doc.core_properties.author:
            metadata["author"] = doc.core_properties.author
        if doc.core_properties.subject:
            metadata["subject"] = doc.core_properties.subject
        if doc.core_properties.keywords:
            metadata["keywords"] = doc.core_properties.keywords

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        text = "\n".join(text_parts)

        return ExtractionResult(
            text=text,
            page_count=len(doc.paragraphs),
            file_type="docx",
            metadata=metadata,
        )
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}") from e


def extract_text_from_buffer(
    buffer: io.BytesIO, filename: str
) -> ExtractionResult:
    """
    Extract text from a buffer based on file extension.

    Auto-detects format based on file extension:
    - .pdf -> PDF extraction
    - .docx -> DOCX extraction
    - .txt -> plain text
    - .md -> markdown (treated as plain text)
    - .yaml/.yml -> YAML (returned as text)
    - .json -> JSON (returned as formatted text)
    - .toml -> TOML (returned as text)
    - .ini/.env -> INI/ENV (returned as text)

    Args:
        buffer: BytesIO buffer containing file content
        filename: Original filename to determine extraction method

    Returns:
        ExtractionResult with extracted text
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return extract_text_from_pdf(buffer)

    if suffix == ".docx":
        return extract_text_from_docx(buffer)

    buffer.seek(0)
    content_bytes = buffer.read()
    content = content_bytes.decode("utf-8", errors="replace")

    text_formats = {".txt", ".md", ".markdown", ".yaml", ".yml", ".json", ".toml", ".ini", ".env", ".cfg", ".conf", ".log"}

    if suffix in text_formats:
        if suffix == ".json":
            try:
                parsed = json.loads(content)
                content = json.dumps(parsed, indent=2, ensure_ascii=False)
            except json.JSONDecodeError:
                pass
        elif suffix in {".yaml", ".yml"}:
            try:
                parsed = yaml.safe_load(content)
                content = yaml.dump(parsed, allow_unicode=True, default_flow_style=False)
            except yaml.YAMLError:
                pass

        return ExtractionResult(
            text=content,
            page_count=1,
            file_type=suffix.lstrip("."),
            metadata={"source_file": filename},
        )

    return ExtractionResult(
        text=content,
        page_count=1,
        file_type="unknown",
        metadata={"source_file": filename},
    )


def download_and_extract(url: str, filename: str) -> ExtractionResult:
    """
    Download a file from URL and extract its text content.

    Args:
        url: URL to download from
        filename: Original filename (used for format detection)

    Returns:
        ExtractionResult with extracted text

    Raises:
        Exception: If download or extraction fails
    """
    buffer = download_file(url)
    return extract_text_from_buffer(buffer, filename)


def extract_text_preview(text: str, max_chars: int = 2000) -> str:
    """
    Extract a preview of text, limiting to max_chars.

    Args:
        text: Full text to preview
        max_chars: Maximum characters to return

    Returns:
        Preview text (truncated if necessary)
    """
    if len(text) <= max_chars:
        return text

    return text[:max_chars] + "..."


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1:
        test_url = sys.argv[1]
        test_filename = sys.argv[2] if len(sys.argv) > 2 else "document.pdf"

        print(f"Downloading {test_filename} from {test_url}...")
        result = download_and_extract(test_url, test_filename)

        print(f"File type: {result.file_type}")
        print(f"Page count: {result.page_count}")
        print(f"Text length: {len(result.text)}")
        print(f"Preview:\n{extract_text_preview(result.text, 500)}")
    else:
        print("Usage: python doc_extractor.py <url> <filename>")
